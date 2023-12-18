import tkinter as tk
from tkinter import ttk, filedialog
from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from PyPDF2 import PdfFileWriter, PdfFileReader
import fitz
import language_tool_python
class DocumentFormatterApp:
    # Инициализация приложения
    def __init__(self, master):
        # Установка главного окна
        self.master = master
        master.title("Форматирование документа")

        # Определение стилей для кнопок
        style = ttk.Style()
        style.theme_use('clam')
        style.configure('Text.TButton', background='#568fba', padding=2, font=('Google Sans', 14), foreground='white')
        style.map('Text.TButton', background=[('active', '#274873')])

        # Центрирование элементов в окне
        master.columnconfigure(0, weight=1)
        master.columnconfigure(1, weight=1)
        master.columnconfigure(2, weight=1)
# Переменные для хранения путей к файлам
        self.file_path = tk.StringVar()
        self.download_path = tk.StringVar()
# Элементы пользовательского интерфейса
        self.file_label = ttk.Label(master, text="Выберите файл:")
        self.file_label.grid(row=0, column=0, padx=10, pady=10)

        self.file_entry = ttk.Entry(master, textvariable=self.file_path, width=30)
        self.file_entry.grid(row=0, column=1, padx=10, pady=10)

        self.browse_button = ttk.Button(master, text="Обзор", command=self.choose_file, style="Text.TButton")
        self.browse_button.grid(row=0, column=2, padx=5, pady=5)

        self.download_folder_label = ttk.Label(master, text="Выберите папку для скачивания:")
        self.download_folder_label.grid(row=1, column=0, padx=10, pady=10)

        self.download_folder_entry = ttk.Entry(master, textvariable=self.download_path, width=30)
        self.download_folder_entry.grid(row=1, column=1, padx=10, pady=10)

        self.download_folder_button = ttk.Button(master, text="Обзор", command=self.choose_download_folder, style="Text.TButton")
        self.download_folder_button.grid(row=1, column=2, padx=5, pady=5)

        self.format_button = ttk.Button(master, text="Отформатировать", command=self.format_document, style="Text.TButton")
        self.format_button.grid(row=2, column=1, padx=10, pady=20)
# LanguageTool для проверки грамматики и правописания
        self.language_tool_ru = language_tool_python.LanguageTool('ru-RU')
        self.language_tool_en = language_tool_python.LanguageTool('en-US')
# Функция выбора файла
    def choose_file(self):
        file_path = filedialog.askopenfilename(filetypes=[("Word Files", "*.docx"), ("PDF Files", "*.pdf")])
        self.file_path.set(file_path)
# Функция выбора папки для скачивания
    def choose_download_folder(self):
        download_folder = filedialog.askdirectory()
        self.download_path.set(download_folder)
# Функция форматирования документа
    def format_document(self):
        file_path = self.file_path.get()
        file_ext = file_path.split('.')[-1].lower()

        if file_ext == "docx":
            self.format_word_document()
        elif file_ext == "pdf":
            self.format_pdf_document()
    def format_word_document(self):
        # Открытие документа Word с использованием библиотеки docx
        doc = Document(self.file_path.get())
# Установка отступов страницы в миллиметрах
        sections = doc.sections
        for section in sections:
            section.top_margin = Cm(2)  # 20 мм
            section.bottom_margin = Cm(2)  # 20 мм
            section.left_margin = Cm(3)  # 30 мм
            section.right_margin = Cm(1.5)  # 15 мм
# Форматирование текста и исправление ошибок
        for paragraph in doc.paragraphs:
            paragraph.style.font.name = 'Times New Roman'
            paragraph.style.font.size = Pt(14)
            paragraph.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.5
            paragraph.paragraph_format.first_line_indent = Inches(0.49)
            #paragraph.paragraph_format.left_indent = Inches(0.49)  # 1,25 см
            paragraph.paragraph_format.space_after = Pt(12)
 # Применение отступа для первой строки
            paragraph.paragraph_format.first_line_indent = Pt(14 * 1.25)
# Проверка грамматики и правописания
            text = paragraph.text
            matches = self.language_tool_ru.check(text)
            matchesii = self.language_tool_en.check(text)
            if matches:
                print(f"Errors in paragraph: {text}")
                for match in matches:
                    print(f"Error: {match.ruleId}, Message: {match.message}")
                    # Автоматическое исправление ошибок правописания
                    corrected_text = self.language_tool_ru.correct(text)
                    paragraph.text = corrected_text
            if matchesii:
                print(f"Errors in paragraph: {text}")
                for match in matchesii:
                    print(f"Error: {match.ruleId}, Message: {match.message}")
                    # Автоматическое исправление ошибок правописания
                    corrected_text = self.language_tool_en.correct(text)
                    paragraph.text = corrected_text
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.name = "Times New Roman"
                            run.font.size = Pt(14)  # Пример значения размера шрифта
                            paragraph.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                            paragraph.paragraph_format.line_spacing = 1.5
                            paragraph.paragraph_format.space_after = Pt(12)
                # Обработка нумерованных списков
        for para in doc.paragraphs:
            if para.style.name.startswith('List Paragraph'):
                for run in para.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(14)
# Поиск и выравнивание изображений по центру
        for shape in doc.inline_shapes:
            if shape.type == 1:
                shape.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        for paragraph in doc.paragraphs:
            for run in paragraph.runs:
                run.font.name = "Times New Roman"
                run.font.size = Pt(14)
            paragraph.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            paragraph.paragraph_format.line_spacing = 1.5
            paragraph.paragraph_format.space_after = Pt(12)
            # Проверяем наличие изображений внутри параграфа
            for run in paragraph.runs:
                if run._element.tag.endswith('drawing') and run._element.drawing:
                    for shape in run._element.drawing:
                        if shape.type == 1:  # InlineShapePicture
                            shape.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # Center
# Сохранение отформатированного текста
        formatted_path = "temp_formatted.docx"
        doc.save(formatted_path)
        self.move_and_cleanup(formatted_path)
    def format_pdf_document(self):
        # Открытие PDF-документа с использованием библиотеки fitz
        doc = fitz.open(self.file_path.get())
        output_doc = fitz.open()
        for page_num in range(doc.page_count):
            page = doc[page_num]
            new_page = output_doc.new_page(width=page.rect.width, height=page.rect.height)

            # Применение отступов страницы в миллиметрах
            new_page.set_margins(left=30.0 / 25.4, right=(page.rect.width - 15.0) / 25.4, top=20.0 / 25.4, bottom=(page.rect.height - 20.0) / 25.4)
# Итерация по изображениям страницы
            for img_index in range(page.get_image_count()):
                img = page.get_image(img_index)
                img_rect = img[0]
                img_width = img_rect.width
                img_height = img_rect.height

                # Вычисление новой позиции для центрирования изображения
                new_x = (page.rect.width - img_width) / 2
                new_y = (page.rect.height - img_height) / 2

                img_rect.x0 = new_x
                img_rect.x1 = new_x + img_width
                img_rect.y0 = new_y
                img_rect.y1 = new_y + img_height

                new_page.insert_image(img_rect, img_index=img_index)
# Сохранение отформатированного PDF
        formatted_path = "temp_formatted.pdf"
        output_doc.save(formatted_path)
        self.move_and_cleanup(formatted_path)
# Функция перемещения и очистки
    def move_and_cleanup(self, formatted_path):
        download_folder = self.download_path.get()
        new_file_path = f"{download_folder}/formatted_document.{formatted_path.split('.')[-1]}"
        import shutil
        shutil.move(formatted_path, new_file_path)